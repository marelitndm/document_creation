import os
import re
import tempfile
import streamlit as st
import google.generativeai as genai
from googleapiclient.discovery import build
from google.oauth2 import service_account

# New imports for Markdown → HTML → DOCX conversion
import markdown
from bs4 import BeautifulSoup, Tag, NavigableString
from docx import Document
from googleapiclient.http import MediaFileUpload

# ------------------------------
# 1. Load API Keys and Credentials from st.secrets
# ------------------------------
GEMINI_API_KEY = st.secrets["gemini"]["GEMINI_API_KEY"]
if not GEMINI_API_KEY:
    st.error("GEMINI_API_KEY not found in st.secrets. Please check your configuration.")
    st.stop()

genai.configure(api_key=GEMINI_API_KEY)

# ------------------------------
# 1a. Obtain DOCX Template from st.secrets (optional)
# ------------------------------
DOCX_TEMPLATE_PATH = st.secrets.get("docx_template_path", None)
# No error is raised if template is not provided; a blank document will be used.

# ------------------------------
# 2. Generation configuration (Gemini)
# ------------------------------
generation_config = {
    "temperature": 0,
    "top_p": 0.95,
    "top_k": 64,
    "max_output_tokens": 65536,
    "response_mime_type": "text/plain",
}

# ------------------------------
# 3. Helper function to upload files to Gemini
# ------------------------------
def upload_to_gemini(file_path, mime_type=None):
    file_obj = genai.upload_file(file_path, mime_type=mime_type)
    st.write(f"Uploaded file '{file_obj.display_name}' as: {file_obj.uri}")
    return file_obj

# ------------------------------
# 4. Document generation function (returns Markdown text)
# ------------------------------
def generate_document(document_type, inputs, voice_file_paths=None):
    if document_type == "Meeting Notes":
        prompt_file = "meetingprompt.txt"
        prompt_text = (
            f"Task: Create a meeting summary note in Markdown format.\n"
            f"Meeting Title: {inputs.get('meeting_title', '')}\n"
            f"Meeting Purpose: {inputs.get('meeting_purpose', '')}\n"
            f"Attendees: {inputs.get('attendees', '')}\n"
            f"Discussion Points: {inputs.get('discussion_points', '')}\n"
            f"Action Items: {inputs.get('action_items', '')}\n"
            f"Outcomes: {inputs.get('outcomes', '')}\n"
        )
        if voice_file_paths:
            prompt_text += "\nAdditionally, please integrate key insights and important points from the attached voice recordings to create a comprehensive meeting summary."
    elif document_type == "Mission Document":
        prompt_file = "missionprompt.txt"
        prompt_text = (
            f"Task: Create a mission document in Markdown format.\n"
            f"Project Title: {inputs.get('project_title', '')}\n"
            f"Project Objective: {inputs.get('project_objective', '')}\n"
            f"Scope and Deliverables: {inputs.get('scope_deliverables', '')}\n"
            f"Stakeholders: {inputs.get('stakeholders', '')}\n"
            f"Ownership and Accountability: {inputs.get('ownership_accountability', '')}\n"
            f"Key Decisions: {inputs.get('key_decisions', '')}\n"
            f"Action Plan: {inputs.get('action_plan', '')}\n"
        )
        if voice_file_paths:
            prompt_text += "\nAdditionally, please integrate key insights from the attached voice recordings into the mission document."
    else:
        prompt_file = None
        prompt_text = "No valid document type selected."

    if prompt_file and os.path.exists(prompt_file):
        with open(prompt_file, "r") as f:
            sop_text = f.read()
    else:
        sop_text = ""

    model_instance = genai.GenerativeModel(
        model_name="gemini-2.0-flash-thinking-exp-01-21",
        generation_config=generation_config,
        system_instruction=sop_text,
    )
    chat_session = model_instance.start_chat(history=[])

    if voice_file_paths:
        file_objs = []
        for voice_file_path in voice_file_paths:
            ext = os.path.splitext(voice_file_path)[1].lower()
            mime = None
            if ext == '.ogg':
                mime = "audio/ogg"
            elif ext == '.mp3':
                mime = "audio/mpeg"
            elif ext == '.wav':
                mime = "audio/wav"
            file_obj = upload_to_gemini(voice_file_path, mime_type=mime)
            file_objs.append(file_obj)
        chat_session = model_instance.start_chat(history=[{"role": "user", "parts": file_objs}])

    response = chat_session.send_message(prompt_text)
    cleaned_response_text = response.text.replace("markdown", "")
    return cleaned_response_text

# ------------------------------
# 5. Helper function: Convert Markdown to HTML and parse with BeautifulSoup
# ------------------------------
def markdown_to_html(md_text):
    html = markdown.markdown(md_text, extensions=['tables'])
    soup = BeautifulSoup(html, "html.parser")
    return str(soup)

# ------------------------------
# 6. Helper function: Convert HTML to a DOCX file using python-docx with a template.
# ------------------------------
def html_to_docx(html, output_path, template_path=None):
    if template_path and os.path.exists(template_path):
        doc = Document(template_path)
    else:
        doc = Document()

    soup = BeautifulSoup(html, 'html.parser')
    elements = soup.body.contents if soup.body else soup.contents

    for elem in elements:
        if isinstance(elem, (str, NavigableString)):
            text = str(elem).strip()
            if text:
                doc.add_paragraph(text)
            continue

        if not isinstance(elem, Tag):
            continue

        tag = elem.name.lower() if elem.name else ""
        if tag in ['h1', 'h2', 'h3', 'h4', 'h5', 'h6']:
            level = int(tag[1])
            doc.add_heading(elem.get_text(strip=True), level=level)
        elif tag == 'p':
            doc.add_paragraph(elem.get_text(strip=True))
        elif tag == 'ul':
            for li in elem.find_all('li'):
                doc.add_paragraph(li.get_text(strip=True), style='List Bullet')
        elif tag == 'ol':
            for li in elem.find_all('li'):
                doc.add_paragraph(li.get_text(strip=True), style='List Number')
        elif tag == 'table':
            rows = elem.find_all('tr')
            if rows:
                first_row = rows[0]
                cols = first_row.find_all(['th', 'td'])
                num_cols = len(cols)
                table = doc.add_table(rows=0, cols=num_cols)
                for row in rows:
                    cells = row.find_all(['th', 'td'])
                    doc_row = table.add_row().cells
                    for idx, cell in enumerate(cells):
                        doc_row[idx].text = cell.get_text(separator=' ').strip()
        else:
            text = elem.get_text(separator=' ', strip=True)
            if text:
                doc.add_paragraph(text)

    doc.save(output_path)
    return output_path

# ------------------------------
# 7. Helper function: Upload a DOCX file to Google Drive
# ------------------------------
def upload_docx_to_drive(docx_path, title, folder_id=None):
    SCOPES = ["https://www.googleapis.com/auth/drive"]
    credentials = service_account.Credentials.from_service_account_info(
        st.secrets["google"],
        scopes=SCOPES
    )
    drive_service = build("drive", "v3", credentials=credentials)

    file_metadata = {
        "name": title,
        "mimeType": "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    }
    if folder_id:
        file_metadata["parents"] = [folder_id]

    media = MediaFileUpload(
        docx_path,
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
    file = drive_service.files().create(
        body=file_metadata,
        media_body=media,
        fields="id"
    ).execute()
    return file.get("id")

# ------------------------------
# 8. Streamlit App UI
# ------------------------------
st.title("Documentation Setup Tool")

document_type = st.sidebar.selectbox("Select Document Type", ["Meeting Notes", "Mission Document"])

if document_type == "Meeting Notes":
    st.header("Generate Meeting Notes")
    meeting_title = st.text_input("Meeting Title", "Meeting Notes")
    meeting_purpose = st.text_area("Meeting Purpose")
    attendees = st.text_area("Attendees (comma-separated)")
    discussion_points = st.text_area("Discussion Points")
    action_items = st.text_area("Action Items")
    outcomes = st.text_area("Outcomes")
    voice_files = st.file_uploader("Optional: Upload Voice Recordings", type=["ogg", "mp3", "wav"], accept_multiple_files=True)

    if st.button("Generate Meeting Notes Document"):
        inputs = {
            "meeting_title": meeting_title,
            "meeting_purpose": meeting_purpose,
            "attendees": attendees,
            "discussion_points": discussion_points,
            "action_items": action_items,
            "outcomes": outcomes,
        }
        voice_file_paths = []
        if voice_files:
            for voice_file in voice_files:
                suffix = os.path.splitext(voice_file.name)[1]
                with tempfile.NamedTemporaryFile(delete=False, suffix=suffix) as tmp:
                    tmp.write(voice_file.read())
                    voice_file_paths.append(tmp.name)
        with st.spinner("Generating document..."):
            result = generate_document("Meeting Notes", inputs, voice_file_paths)
        st.success("Document generated!")
        st.text_area("Generated Markdown", result, height=400)

        with st.spinner("Converting Markdown to HTML..."):
            html_content = markdown_to_html(result)
        with st.spinner("Converting HTML to DOCX using shared template..."):
            with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp_docx:
                docx_path = tmp_docx.name
            html_to_docx(html_content, docx_path, template_path=DOCX_TEMPLATE_PATH)
        with st.spinner("Uploading DOCX to Google Drive..."):
            doc_id = upload_docx_to_drive(docx_path, title=meeting_title, folder_id=st.secrets["google"].get("drive_folder_id"))
        st.success(f"Document saved to Google Drive! Document ID: {doc_id}")
        st.write("Open your document using this URL:")
        st.write(f"https://drive.google.com/file/d/{doc_id}/view")

elif document_type == "Mission Document":
    st.header("Generate Mission Document")
    st.write("**Debug:** Mission Document mode active.")
    project_title = st.text_input("Project Title", "Mission Document")
    project_objective = st.text_area("Project Objective")
    scope_deliverables = st.text_area("Scope and Deliverables")
    stakeholders = st.text_area("Stakeholders (names and roles)")
    ownership_accountability = st.text_area("Ownership and Accountability")
    key_decisions = st.text_area("Key Decisions")
    action_plan = st.text_area("Action Plan")
    voice_files = st.file_uploader("Optional: Upload Voice Recordings", type=["ogg", "mp3", "wav"], accept_multiple_files=True)

    if st.button("Generate Mission Document"):
        inputs = {
            "project_title": project_title,
            "project_objective": project_objective,
            "scope_deliverables": scope_deliverables,
            "stakeholders": stakeholders,
            "ownership_accountability": ownership_accountability,
            "key_decisions": key_decisions,
            "action_plan": action_plan,
        }
        voice_file_paths = []
        if voice_files:
            for voice_file in voice_files:
                suffix = os.path.splitext(voice_file.name)[1]
                with tempfile.NamedTemporaryFile(delete=False, suffix=suffix) as tmp:
                    tmp.write(voice_file.read())
                    voice_file_paths.append(tmp.name)
        with st.spinner("Generating document..."):
            result = generate_document("Mission Document", inputs, voice_file_paths)
        st.success("Document generated!")
        st.text_area("Generated Markdown", result, height=400)

        with st.spinner("Converting Markdown to HTML..."):
            html_content = markdown_to_html(result)
        with st.spinner("Converting HTML to DOCX using shared template..."):
            with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp_docx:
                docx_path = tmp_docx.name
            html_to_docx(html_content, docx_path, template_path=DOCX_TEMPLATE_PATH)
        with st.spinner("Uploading DOCX to Google Drive..."):
            doc_id = upload_docx_to_drive(docx_path, title=project_title, folder_id=st.secrets["google"].get("drive_folder_id"))
        st.success(f"Document saved to Google Drive! Document ID: {doc_id}")
        st.write("Open your document using this URL:")
        st.write(f"https://drive.google.com/file/d/{doc_id}/view")
